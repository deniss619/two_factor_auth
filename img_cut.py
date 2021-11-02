from PIL import Image
# img = Image.open('zzz.png')
# area = (5*38+10, 2, 5*38+10+38, 40)
# cropped_img = img.crop(area)
# cropped_img.show()
# cropped_img.save("1.png")
# name=0
# img.show()
# for i in range(20):
#     #img = Image.open("Emojis.png")
#     for j in range(20):
#         name+=1
#         if j==0:
#             g=0
#         else:
#             g=2
#         area = (38*j, ((38*i)), 38*j+38, 38*i+38)
#         cropped_img = img.crop(area)
#         #cropped_img.show()
#         cropped_img.save(str(name)+'.png')



import numpy as np
import pyautogui
import imutils
import cv2
import docx


def add_to_word(text, organization):
    try:
        doc = docx.Document("отчет.docx")
    except:
        doc = docx.Document()
        doc.add_paragraph(text)
    for paragraph in doc.paragraphs:
        if paragraph.text != text:
            text_in_docx = 0
        else:
            text_in_docx = 1
            break
    if not text_in_docx:
        doc.add_paragraph(text)
    doc.add_picture('screenshot.png', width=docx.shared.Cm(10))
    doc.save(f'Отчет {organization}.docx')

organization = 'name'
for i in range(3):
    image = pyautogui.screenshot()
    image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    cv2.imwrite('screenshot.png', image)
    add_to_word('По счету 68.04.1 найдены пустые аналитики в журнале проводок', organization)






